# File-level suppression removed per audit (V143 hardening).
# Per-line justified suppressions (e.g., '# noqa: S3776 ...') are preserved.
"""
word_parser.py — FireAI Word Specification Parser
Parses project specifications from Word documents.
"""

import logging
import re
from dataclasses import dataclass, field
from typing import Dict, List

from parsers._base import ParserBase
from parsers._path_security import UnsafePathError

logger = logging.getLogger("fireai.word_parser")


@dataclass
class WordParseResult:
    """Result of parsing Word document."""

    source_file: str
    success: bool
    title: str = ""
    project_name: str = ""
    floor: str = ""
    building: str = ""
    ceiling_specs: List[Dict] = field(default_factory=list)
    requirements: List[str] = field(default_factory=list)
    notes: List[str] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)


class WordParser(ParserBase):
    """
    Parses Word documents for project specifications.
    """

    allowed_extensions = {'.docx'}
    max_file_size_bytes = int(os.getenv("FIREAI_WORD_MAX_FILE_SIZE_BYTES", 25 * 1024 * 1024))

    FLOOR_PATTERNS = [
        r'floor\s*(\d+)',
        r'level\s*(\d+)',
        r'Floor\s*(\d+)',
        r'Level\s*(\d+)',
        r'(\d+)st\s*floor',
        r'(\d+)nd\s*floor',
        r'(\d+)rd\s*floor',
        r'(\d+)th\s*floor',
    ]

    BUILDING_PATTERNS = [
        r'building\s*([A-Z])',
        r'tower\s*([A-Z])',
        r'block\s*([A-Z])',
        r'([A-Z])\s*building',
    ]

    CEILING_PATTERNS = [
        r'ceiling\s*height[:\s]*(\d+\.?\d*)\s*m',
        r'height[:\s]*(\d+\.?\d*)\s*m',
        r'flat\s*ceiling[:\s]*(\d+\.?\d*)',
        r'suspended\s*ceiling[:\s]*(\d+\.?\d*)',
    ]

    def __init__(self):
        pass

    def parse(self, file_path: str) -> WordParseResult:
        try:
            safe_path = self.validate_input(file_path)
        except FileNotFoundError as e:
            return WordParseResult(source_file=file_path, success=False, errors=[str(e)])
        except UnsafePathError as e:
            return WordParseResult(source_file=file_path, success=False, errors=[f"SECURITY: {e}"])

        file_path = str(safe_path)
        result = WordParseResult(source_file=file_path, success=False)

        try:
            from docx import Document

            doc = Document(str(safe_path))

            all_text = '\n'.join(p.text for p in doc.paragraphs)

            result.title = self._extract_title(doc.paragraphs)
            result.project_name = self._extract_project_name(all_text)
            result.floor = self._extract_floor(all_text)
            result.building = self._extract_building(all_text)
            result.ceiling_specs = self._extract_ceiling_specs(all_text)
            result.requirements = self._extract_requirements(doc.paragraphs)
            result.notes = self._extract_notes(doc.paragraphs)

            result.success = bool(result.title or result.project_name or result.floor)

        except ImportError:
            result.errors.append("Missing dependency: python-docx not installed")
        except Exception as e:
            result.errors.append(f"Parse error: {type(e).__name__}: {e}")

        return result

    def _extract_title(self, paragraphs) -> str:
        for para in paragraphs:
            if para.style.name.startswith('Heading'):
                return para.text.strip()
        return ""

    def _extract_project_name(self, text: str) -> str:
        patterns = [
            r'Project[:\s]*([^\n]+)',
            r'Building[:\s]*([^\n]+)',
            r'Tower\s*([A-Z])',
        ]

        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()

        return ""

    def _extract_floor(self, text: str) -> str:
        for pattern in self.FLOOR_PATTERNS:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return f"Floor {match.group(1)}"
        return ""

    def _extract_building(self, text: str) -> str:
        for pattern in self.BUILDING_PATTERNS:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return f"Building {match.group(1)}"
        return ""

    def _extract_ceiling_specs(self, text: str) -> List[Dict]:
        specs = []

        for pattern in self.CEILING_PATTERNS:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    height = float(match.group(1))
                    specs.append({
                        'type': 'flat',
                        'height_m': height,
                    })
                except (ValueError, IndexError):
                    continue

        return specs

    def _extract_requirements(self, paragraphs) -> List[str]:
        requirements = []

        for para in paragraphs:
            text = para.text.strip()

            if text.startswith('•') or text.startswith('- ') or text.startswith('* '):
                clean_text = text.lstrip('•-* ').strip()

                if any(kw in clean_text.lower() for kw in [
                    'detector', 'alarm', 'fire', 'sprinkler',
                    'system', 'zone', 'coverage', 'code'
                ]):
                    requirements.append(clean_text)

        return requirements

    def _extract_notes(self, paragraphs) -> List[str]:
        notes = []
        in_notes = False

        for para in paragraphs:
            text = para.text.strip()

            if 'note' in text.lower() and len(text) < 20:
                in_notes = True
                continue

            if in_notes and text:
                notes.append(text)

        return notes[:10]


def parse_word(file_path: str) -> WordParseResult:
    """Quick parse Word file."""
    parser = WordParser()
    return parser.parse(file_path)
